"""
generate_mechanistic_report.py
==============================
Builds the Word write-up for the mechanistic crowded predator-prey model
(crowded_mechanistic.py). Written so a project teammate who has not seen how the
model was built can read it top to bottom and understand exactly what equations
we propose and where every term comes from.

Formatting notes:
- Genuine equations are inserted as NATIVE, editable Word math
  (LaTeX -> MathML -> OMML via Office's MML2OMML.XSL).
- Subscripts in ordinary text/tables/schemes are REAL subscripts (run.subscript),
  written here with a _{...} (and ^{...}) mini-markup that add_rich() parses.
- No em dashes anywhere.
Run after crowded_mechanistic.py (it consumes the two figures).
"""

import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import latex2mathml.converter as _l2m
from lxml import etree

HERE = os.path.dirname(os.path.abspath(__file__))
NAVY = RGBColor(0x1F, 0x3A, 0x5F)
AMBER = RGBColor(0xC4, 0x78, 0x28)
GREY = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
MONO, SERIF = "Consolas", "Georgia"

# --- native Word equations (LaTeX -> MathML -> OMML) -------------------------
_XSL_CANDIDATES = [
    r"C:/Program Files/Microsoft Office/root/Office16/MML2OMML.XSL",
    r"C:/Program Files (x86)/Microsoft Office/root/Office16/MML2OMML.XSL",
    r"C:/Program Files/Microsoft Office/Office16/MML2OMML.XSL",
]
_XSL_PATH = next((p for p in _XSL_CANDIDATES if os.path.exists(p)), None)
_OMML_XSLT = etree.XSLT(etree.parse(_XSL_PATH)) if _XSL_PATH else None


def _latex_to_omml(latex):
    return _OMML_XSLT(etree.fromstring(_l2m.convert(latex))).getroot()


doc = Document()
doc.styles["Normal"].font.name = SERIF
doc.styles["Normal"].font.size = Pt(10.5)

# --- rich text with real subscripts/superscripts -----------------------------
_TOKEN = re.compile(r'(_\{[^}]*\}|\^\{[^}]*\})')


def add_rich(p, text, size=10.5, color=None, italic=False, bold=False, name=SERIF):
    """Add runs to paragraph p, turning _{..} into real subscripts and
    ^{..} into real superscripts."""
    for part in _TOKEN.split(text):
        if not part:
            continue
        sub = sup = False
        if part.startswith('_{') and part.endswith('}'):
            part, sub = part[2:-1], True
        elif part.startswith('^{') and part.endswith('}'):
            part, sup = part[2:-1], True
        r = p.add_run(part)
        r.font.size = Pt(size)
        r.font.name = name
        r.italic = italic
        r.bold = bold
        if color is not None:
            r.font.color.rgb = color
        if sub:
            r.font.subscript = True
        if sup:
            r.font.superscript = True
    return p


def _shade(p, fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)


def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY if level <= 2 else AMBER
        run.font.name = SERIF
    return h


def body(text, italic=False, size=10.5, color=None, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    add_rich(p, text, size=size, color=color, italic=italic)
    return p


def omath(latex, center=True):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    if _OMML_XSLT is not None:
        p._p.append(_latex_to_omml(latex))
    else:
        add_rich(p, latex, name=MONO, size=9.5)
    return p


def scheme(label, text):
    """A reaction line: bold label, then the scheme with real subscripts."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    add_rich(p, label, bold=True, color=NAVY)
    add_rich(p, "   " + text)
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        add_rich(p, bold_lead, bold=True)
    add_rich(p, text)
    return p


def caveat(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    add_rich(p, "Note: " + text, italic=True, size=9.5, color=AMBER)
    return p


def make_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = ""
        add_rich(c.paragraphs[0], h, bold=True, size=9)
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = ""
            add_rich(cells[j].paragraphs[0], str(val), size=8.5)
    if widths:
        for j, w in enumerate(widths):
            for r in t.rows:
                r.cells[j].width = Inches(w)
    return t


# =============================================================================
# TITLE
# =============================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("A Mechanistic Model of Molecular Crowding in the\n"
                  "DNA Predator and Prey Oscillator")
r.bold = True
r.font.size = Pt(19)
r.font.name = SERIF
r.font.color.rgb = NAVY
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("How crowding reshapes the Fujii and Rondelez PEN oscillator: "
                "the model we propose, and how it was built.")
r.italic = True
r.font.size = Pt(12)
r.font.name = SERIF
r.font.color.rgb = GREY
m = doc.add_paragraph()
m.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = m.add_run("BIOMOD dry lab. Written for a project teammate. "
              "Implementation: crowded_mechanistic.py")
r.font.size = Pt(10)
r.font.color.rgb = GREY
doc.add_paragraph()

# =============================================================================
# 1. WHAT THIS DOCUMENT IS
# =============================================================================
heading("1.  What this document is", 1)
body("This describes the model our team built for the dry lab part of the "
     "project: a model of how molecular crowding (adding an inert polymer such "
     "as PEG or Ficoll to the tube) changes the DNA predator and prey "
     "oscillator. It is written for a teammate who knows the biology at a high "
     "level but has not seen how this model was put together. By the end you "
     "will know exactly which equations we propose, what every symbol means, "
     "and why each piece is there.")
body("The structure is: Section 2 states the final model in full. Sections 3 to "
     "8 derive it, starting from the published dilute oscillator and adding the "
     "crowding physics one layer at a time. Sections 9 to 11 give the "
     "predictions, the optional simpler version, and what the wet lab should "
     "measure to test it.")

# =============================================================================
# 2. THE FINAL MODEL  (banner box)
# =============================================================================
banner = doc.add_paragraph()
_shade(banner, "1F3A5F")
banner.paragraph_format.space_before = Pt(6)
banner.paragraph_format.space_after = Pt(2)
rb = banner.add_run("  2.  THE FINAL MODEL (THE EQUATIONS WE PROPOSE)")
rb.bold = True
rb.font.size = Pt(13)
rb.font.name = SERIF
rb.font.color.rgb = WHITE

body("This is the complete model. If you read only one section, read this one. "
     "Everything after it explains where each piece comes from.")

body("The model tracks seven concentrations in nanomolar (nM): free prey "
     "strand N, free predator strand P, free template G, and four "
     "intermediate complexes (two molecules bound together): "
     "C_{NG} (prey on template), C_{NP} (prey on predator), "
     "C_{NE} (prey on the degrading enzyme), and C_{PE} (predator on the "
     "degrading enzyme). The free enzyme is whatever is not currently bound:")
omath(r"E = rec - C_{NE} - C_{PE}")

body("The seven concentrations change in time according to these equations. "
     "Each k is a rate constant (defined just below); the crowding-sensitive "
     "ones are highlighted in the next box.")
omath(r"\frac{dN}{dt} = -k_{on}^{NG}NG + k_{off}^{NG}C_{NG} + 2k_{cat}^{g}C_{NG} "
      r"- k_{on}^{NP}NP + k_{off}^{NP}C_{NP} - k_{on}^{eN}NE + k_{off}^{eN}C_{NE} + \lambda_N")
omath(r"\frac{dP}{dt} = -k_{on}^{NP}NP + k_{off}^{NP}C_{NP} + 2k_{cat}^{p}C_{NP} "
      r"- k_{on}^{eP}PE + k_{off}^{eP}C_{PE} + \lambda_P")
omath(r"\frac{dG}{dt} = -k_{on}^{NG}NG + k_{off}^{NG}C_{NG} + k_{cat}^{g}C_{NG}")
omath(r"\frac{dC_{NG}}{dt} = k_{on}^{NG}NG - k_{off}^{NG}C_{NG} - k_{cat}^{g}C_{NG}")
omath(r"\frac{dC_{NP}}{dt} = k_{on}^{NP}NP - k_{off}^{NP}C_{NP} - k_{cat}^{p}C_{NP}")
omath(r"\frac{dC_{NE}}{dt} = k_{on}^{eN}NE - k_{off}^{eN}C_{NE} - k_{cat}^{eN}C_{NE}")
omath(r"\frac{dC_{PE}}{dt} = k_{on}^{eP}PE - k_{off}^{eP}C_{PE} - k_{cat}^{eP}C_{PE}")

body("Reading one term: in the prey equation, -k_{on}^{NG}NG is prey binding "
     "the template (it leaves the free pool), +k_{off}^{NG}C_{NG} is that "
     "complex falling back apart, and +2k_{cat}^{g}C_{NG} is the catalytic "
     "step that copies the prey, giving two strands where there was one. "
     "lambda_N and lambda_P are a tiny constant background synthesis ('leak') "
     "that keeps the rhythm robust (Section 8).")

# the crowding rules (second shaded line)
banner2 = doc.add_paragraph()
_shade(banner2, "C47828")
banner2.paragraph_format.space_before = Pt(8)
banner2.paragraph_format.space_after = Pt(2)
rb2 = banner2.add_run("  How crowding enters: one number, the crowder volume "
                      "fraction phi")
rb2.bold = True
rb2.font.size = Pt(11.5)
rb2.font.name = SERIF
rb2.font.color.rgb = WHITE

body("Crowding is summarised by the volume fraction phi that the crowder "
     "occupies, set by its weight concentration C and specific volume v-bar:")
omath(r"\varphi = \bar{v}\,\frac{C}{100}")
body("Every rate constant above is then the dilute value times physics that "
     "depends only on phi, the crowder size, and the molecules' sizes. For each "
     "binding step, the binding strength (equilibrium constant K) and the "
     "on-rate change as:")
omath(r"K(\varphi) = K^{0}\,\exp\!\left(-\frac{\Delta\Delta G_{EV}}{RT}\right), "
      r"\quad k_{on}(\varphi)=\frac{k_{on}^{0}}{\eta_{eff}}, "
      r"\quad k_{off}=\frac{k_{on}}{K}")
body("where the excluded-volume free energy that strengthens binding is")
omath(r"\frac{\Delta\Delta G_{EV}}{RT} = \mu_{ex}(r_{AB}) - \mu_{ex}(r_A) - \mu_{ex}(r_B)")
body("and the catalytic constants shift by enzyme class (polymerase and nickase "
     "up, exonuclease down):")
omath(r"k_{cat}^{g},\,k_{cat}^{p} \;\to\; k_{cat}\,e^{+c_{P}\varphi}, "
      r"\qquad k_{cat}^{eN},\,k_{cat}^{eP} \;\to\; k_{cat}\,e^{-c_{E}\varphi}")
body("mu_ex and eta_eff are defined in Section 6. At zero crowder (phi = 0) "
     "every rate constant returns to its dilute value (Table in Section 5) and "
     "the model reproduces the measured oscillation: period about 90 minutes, "
     "predator swinging from about 22 up to 154 nM. That is the whole model.")

# =============================================================================
# 3. BACKGROUND: THE DILUTE OSCILLATOR
# =============================================================================
heading("3.  Background: the oscillator we start from", 1)
body("The system (Fujii and Rondelez, 2013) is three short DNA species and "
     "three enzymes in a buffer. Prey N copies itself on a template G with the "
     "help of a polymerase and a nicking enzyme. Predator P copies itself using "
     "prey N as its template, consuming prey in the process. A single "
     "exonuclease (ttRecJ) degrades both prey and predator. Because the "
     "predator lags the prey, the two populations chase each other up and down "
     "in a sustained oscillation, exactly like foxes and rabbits.")
body("The published description, after simplification, is two equations for the "
     "prey and predator concentrations. We use these as our dilute baseline:")
omath(r"\frac{dN}{dt}=\frac{k_1\,pol\,G\,N}{1+bGN} - k_2\,pol\,N\,P "
      r"- \frac{rec\,k_N\,N}{1+P/K_{m,P}}")
omath(r"\frac{dP}{dt}=k_2\,pol\,N\,P - \frac{rec\,k_P\,P}{1+P/K_{m,P}}")
body("Here pol is the polymerase level and rec is the exonuclease level. The "
     "six constants k_{1}, b, k_{2}, k_{N}, k_{P}, K_{m,P} are 'lumped': each "
     "one secretly bundles several physical steps together. That bundling is "
     "the key point for what follows.")

# =============================================================================
# 4. THE APPROACH
# =============================================================================
heading("4.  The approach: act on the chemistry, not on the published constants", 1)
body("The tempting shortcut is to keep the two equations above and just "
     "multiply each lumped constant by some crowding factor. We deliberately do "
     "not do that, for a concrete reason: crowding does not act on a lumped "
     "constant. It acts on the separate physical steps hidden inside it, and "
     "those steps respond to crowding in opposite directions.")
body("Take k_{1} (prey growth). It hides a binding step (prey finding its "
     "template), a catalytic step (the polymerase copying it), and a diffusion "
     "step (the two molecules meeting). Crowding speeds binding up (by excluded "
     "volume) but slows diffusion down (by viscosity), and separately changes "
     "the enzyme's turnover. A single fitted factor on k_{1} cannot represent "
     "three effects pulling different ways.")
body("So instead we expand the lumped model back into the elementary reactions "
     "it came from (Section 5), and let crowding act on each elementary step "
     "through shared physics (Section 6). The payoff: the same physics that "
     "stabilises a DNA duplex also tightens enzyme binding, and the same "
     "viscosity that slows the polymerase slows the exonuclease. Nothing is fit "
     "step by step; it all follows from phi, the crowder size, and the "
     "molecules' sizes. This is also what lets two crowders of the same volume "
     "fraction (PEG versus Ficoll) give genuinely different answers.")

# =============================================================================
# 5. THE REACTION NETWORK
# =============================================================================
heading("5.  The reaction network", 1)
body("These are the elementary reactions. A double arrow means reversible "
     "binding (on-rate forward, off-rate back); a single arrow is the "
     "irreversible catalytic step. Polymerase and nickase are plentiful here, "
     "so they are folded into the catalytic constants rather than tracked "
     "separately.")
scheme("Prey growth:",
       "N + G  ⇌  C_{NG}  →  G + 2 N    "
       "(on k_{on}^{NG}, off k_{off}^{NG}, copy k_{cat}^{g})")
scheme("Predation:",
       "N + P  ⇌  C_{NP}  →  2 P    "
       "(on k_{on}^{NP}, off k_{off}^{NP}, copy k_{cat}^{p})")
scheme("Prey decay:",
       "N + E  ⇌  C_{NE}  →  E    "
       "(on k_{on}^{eN}, off k_{off}^{eN}, cut k_{cat}^{eN})")
scheme("Predator decay:",
       "P + E  ⇌  C_{PE}  →  E    "
       "(on k_{on}^{eP}, off k_{off}^{eP}, cut k_{cat}^{eP})")
scheme("Conserved totals:",
       "G_{tot} = G + C_{NG} ,    rec = E + C_{NE} + C_{PE}")
body("Two features of the published model now appear on their own, instead of "
     "being written in by hand:")
bullet("the saturating growth term 1/(1 + b G N) is simply finite template: "
       "once the template G_{tot} is fully occupied, growth cannot speed up. ",
       "Template saturation.  ")
bullet("the competition term 1/(1 + P/K_{m,P}) is the shared exonuclease being "
       "a limited resource. The predator binds it tightly (high affinity) and "
       "the prey binds it weakly. When predator is abundant it occupies the "
       "enzyme and shields the prey from being degraded. That self-shielding is "
       "the engine of the oscillation, and here it falls out of conservation "
       "with no extra assumption.", "The oscillation engine.  ")
body("Mathematically, applying a quasi-steady-state assumption to the four "
     "complexes turns these reactions back into the two-variable equations of "
     "Section 3, which fixes the dilute values of every elementary constant. "
     "Those dilute values are:")
make_table(
    ["Lumped constant", "Built from elementary steps", "Dilute value"],
    [["k_{1} (prey growth)", "k_{cat}^{g} times K_{1} (K_{1} = N,G binding)", "2.0e-3"],
     ["b (growth saturation)", "K_{1} / G_{tot}", "4.8e-5"],
     ["k_{2} (predation)", "k_{cat}^{p} times K_{2} (K_{2} = N,P binding)", "3.1e-3"],
     ["k_{N} (prey decay)", "k_{cat}^{eN} / K_{m}^{N} (weak binding, K_{m}^{N} large)", "2.1e-2"],
     ["k_{P} (predator decay)", "k_{cat}^{eP} / K_{m}^{P}", "4.7e-3"],
     ["K_{m,P} (competition)", "(k_{off}^{eP} + k_{cat}^{eP}) / k_{on}^{eP}", "34 nM"]],
    widths=[1.7, 3.2, 1.4])

# =============================================================================
# 6. THE CROWDING PHYSICS
# =============================================================================
heading("6.  The crowding physics (where the crowding rules come from)", 1)

heading("6.1  Excluded volume: why crowding strengthens binding", 2)
body("A crowder molecule physically occupies space, so a reacting molecule has "
     "less room. The free-energy cost of inserting a molecule of radius r into "
     "a bath of crowders (radius R_{c}) packed at fraction phi is given by "
     "scaled-particle theory (Minton; Lebowitz):")
omath(r"\frac{\mu_{ex}(r)}{RT} = -\ln(1-\varphi) + "
      r"\left(3z + 3z^{2} + z^{3}\right)\frac{\varphi}{1-\varphi}, \quad z=\frac{r}{R_c}")
body("When two molecules A and B bind into a more compact product AB, the "
     "product takes up less room than the two separately, so binding 'gives "
     "back' space to the crowders. That is favourable, so crowding pushes "
     "binding forward. The size of that push is the difference in insertion "
     "cost (the equation used in Section 2):")
omath(r"\frac{\Delta\Delta G_{EV}}{RT} = \mu_{ex}(r_{AB}) - \mu_{ex}(r_A) - \mu_{ex}(r_B)")
body("This one function supplies the excluded-volume effect for every binding "
     "step in the network, strand pairing and enzyme binding alike, differing "
     "only by the molecules' radii. The crowder radius R_{c} enters through "
     "z = r/R_{c}, which is exactly why a small crowder (PEG) and a large one "
     "(Ficoll) at the same phi give different results: smaller R_{c} means a "
     "larger excluded-volume push per unit phi.")

heading("6.2  Viscosity: why crowding slows the on-rates", 2)
body("Crowding also makes the solution thicker, so molecules diffuse and meet "
     "more slowly. For a polymer mesh (PEG) the effective viscosity felt by a "
     "molecule depends on its size relative to the mesh spacing xi (Holyst):")
omath(r"\frac{\eta_{eff}(r)}{\eta_0} = \exp\left[b\left(\frac{R_{eff}}{\xi}"
      r"\right)^{a}\right], \quad \frac{1}{R_{eff}^{2}}=\frac{1}{R_h^{2}}+"
      r"\frac{1}{r^{2}}, \quad \xi = R_c\left(\frac{c}{c^{*}}\right)^{-0.75}")
body("A binding on-rate then slows in proportion to 1/eta_eff at the size of "
     "the larger partner (the slower mover sets the meeting rate). Small "
     "strands barely notice; large enzymes feel it strongly. A compact sphere "
     "like Ficoll has no mesh, so it uses the hard-sphere (Mooney) law instead:")
omath(r"\frac{\eta}{\eta_0}=\exp\left[\frac{[\eta]\,\varphi}{1-\varphi/\varphi_{max}}\right]")

heading("6.3  Smaller effects (kept secondary)", 2)
body("Crowding also lowers the activity of water and can sequester buffer "
     "ions, both of which nudge DNA pairing. For the short strands here these "
     "are modest corrections, so they are included but kept small; the ion "
     "channel is switched off by default because our buffer uses magnesium, "
     "whereas the available data are for sodium.")

# =============================================================================
# 7. HOW CROWDING MOVES EACH CONSTANT
# =============================================================================
heading("7.  Putting it together: how crowding moves each rate constant", 1)
body("Each elementary constant now changes with phi in a fixed way. For binding "
     "steps, the strength changes by the excluded-volume free energy, the "
     "on-rate slows by viscosity, and the off-rate is set so the two stay "
     "consistent:")
omath(r"K(\varphi) = K^{0}\exp\!\left(-\frac{\Delta\Delta G_{EV}}{RT}\right), "
      r"\quad k_{on}(\varphi)=\frac{k_{on}^{0}}{\eta_{eff}}, "
      r"\quad k_{off}=\frac{k_{on}}{K}")
body("The catalytic step is intrinsic enzyme chemistry, which excluded volume "
     "cannot predict, so it is taken from enzyme measurements: crowding raises "
     "polymerase and nickase turnover and lowers exonuclease turnover.")
omath(r"k_{cat}^{pol/nick}(\varphi)=k_{cat}\,e^{+c_{P}\varphi}, \qquad "
      r"k_{cat}^{exo}(\varphi)=k_{cat}\,e^{-c_{E}\varphi}")
body("The net effect on the lumped constants, which the code reads off "
     "automatically, is that for PEG the growth and predation constants "
     "(k_{1}, k_{2}) rise while the decay constants (k_{N}, k_{P}) fall, "
     "because the exonuclease is being slowed. The slowing of the exonuclease "
     "is the single most important consequence, since the oscillation lives or "
     "dies on it.")

# =============================================================================
# 8. CALIBRATION
# =============================================================================
heading("8.  Calibration: two small adjustments", 1)
body("Two adjustments turn the elementary constants into a model whose dilute "
     "behaviour matches the experiment.")
bullet("a single factor of 0.8 multiplies the catalytic constants. The quick "
       "estimate of those constants ignores the prey and predator that sit "
       "temporarily inside the complexes; tracking that bound pool slightly "
       "slows the cycle, and this factor restores the measured period.",
       "Timescale (0.8).  ")
bullet("a tiny constant background synthesis (lambda_N about 0.02, lambda_P "
       "about 0.006 nM per minute), a documented feature of this enzyme "
       "toolbox, lifts the populations off zero. Without it the predator can "
       "fall to exactly zero and never recover; with it the oscillation is "
       "robust, as seen in experiment.", "Leak.  ")
body("Result at zero crowder: period 90 minutes, predator swinging from about "
     "22 up to 154 nM, a strong steady oscillation, matching the published "
     "behaviour.")

# =============================================================================
# 9. PREDICTIONS
# =============================================================================
heading("9.  Predictions", 1)
fig1 = os.path.join(HERE, "mechanistic_timeseries.png")
if os.path.exists(fig1):
    doc.add_picture(fig1, width=Inches(6.4))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
body("Figure 1. Settled dynamics at zero crowder and at 5 percent and 10 "
     "percent PEG 8000. The swings shrink as crowding increases.",
     italic=True, size=9)
fig2 = os.path.join(HERE, "mechanistic_predictions.png")
if os.path.exists(fig2):
    doc.add_picture(fig2, width=Inches(6.4))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
body("Figure 2. Period (left) and predator wave size (right) versus crowder "
     "concentration. PEG 8000 drives the wave size to zero near 10 to 12 "
     "percent (the oscillation stops); Ficoll 400 keeps oscillating throughout.",
     italic=True, size=9)
make_table(
    ["Condition", "Period (min)", "Predator wave (nM)", "Oscillating?"],
    [["Dilute (no crowder)", "90", "132", "yes, strong"],
     ["5 percent PEG 8000", "78", "76", "yes"],
     ["8 percent PEG 8000", "73", "19", "yes, near the edge"],
     ["About 11 percent PEG 8000", "stops", "about 0", "no, waves switch off"],
     ["16 percent Ficoll 400", "69", "38", "yes, still going"]],
    widths=[2.1, 1.3, 1.6, 1.7])
body("The three testable claims for the wet lab: (1) crowding shortens the "
     "period; (2) PEG 8000 switches the oscillation off near 10 to 12 percent; "
     "(3) at the same volume fraction, Ficoll 400 is far milder and keeps the "
     "oscillation alive. The third claim, the same packing giving opposite "
     "outcomes for two crowders, is the novel and falsifiable result.")

# =============================================================================
# 10. SIMPLER VERSION
# =============================================================================
heading("10.  If you want a simpler version", 1)
body("If the full seven-equation model is more than an analysis needs, it can "
     "be collapsed to the two-variable form of Section 3 by assuming binding is "
     "fast, with the crowding-dependent constants given by the rules above, "
     "for example:")
omath(r"k_1(\varphi)=\frac{k_{cat}^{g}(\varphi)K_1(\varphi)}{pol}, \quad "
      r"b(\varphi)=\frac{K_1(\varphi)}{G_{tot}}, \quad "
      r"k_2(\varphi)=\frac{k_{cat}^{p}(\varphi)K_2(\varphi)}{pol}")
caveat("Important. This two-variable shortcut is good for intuition and for the "
       "form of each constant's crowding dependence, but it is not quantitatively "
       "faithful here. On its own it mispredicts even the dilute period (about "
       "45 instead of 90 minutes) and loses the oscillation almost immediately, "
       "because it cannot hold the prey and predator that are temporarily bound "
       "inside the complexes. All quantitative predictions (period, the switch-off "
       "concentration, the PEG versus Ficoll difference) come from the full "
       "seven-species model. The interactive explorer runs that full model.")

# =============================================================================
# 11. ASSUMPTIONS AND WHAT TO MEASURE
# =============================================================================
heading("11.  Assumptions and what to measure", 1)
bullet("polymerase and nickase are plentiful and folded into the catalytic "
       "constants; only the exonuclease is treated as a limited, saturating "
       "resource.", "Enzymes.  ")
bullet("fast intermediate folding states are not tracked (assumed quick).",
       "Intermediates.  ")
bullet("the strength of the exonuclease slowdown is taken from a related enzyme "
       "(Exonuclease I), not ttRecJ itself.", "Enzyme data.  ")
bullet("the excluded-volume radii, the viscosity constants, the catalytic "
       "slopes, and the leak size are all literature-anchored but tunable. The "
       "qualitative story (period shortens, PEG switches off, Ficoll milder) is "
       "robust to their exact values; the precise switch-off concentration is "
       "not, so it should be fitted.", "Tunable inputs.  ")
bullet("all Ficoll 400 numbers are extrapolations (no direct Ficoll 400 data "
       "in the surveyed papers).", "Known gaps.  ")
body("Highest-value measurements, in order: (1) ttRecJ activity versus PEG 8000 "
     "and versus Ficoll 400; (2) prey and predator pairing strength versus "
     "crowder; (3) solution viscosity versus crowder at the reaction "
     "temperature; (4) the background leak rate.")

# =============================================================================
# 12. REFERENCES
# =============================================================================
heading("12.  Key references", 1)
body("Access tags: [full] full text seen, [abs] abstract only, [2] via a "
     "secondary source.", italic=True, size=9)
refs = [
    "Fujii and Rondelez (2013). Predator-prey molecular ecosystems. ACS Nano "
    "7:27. Baseline network and the two-variable model. [full]",
    "Montagne et al. (2011). Programming an in vitro DNA oscillator. Mol Syst "
    "Biol 7:466. The PEN toolbox and the leak reactions. [full]",
    "Holyst et al. (2009). Scaling form of viscosity at all length scales. PCCP "
    "11:9025. The scale-dependent viscosity law. [abs]",
    "Knowles et al. (2011). Separating excluded volume from preferential "
    "interactions. PNAS 108:12699. Excluded-volume stabilisation of duplexes. [full]",
    "Zhou, Rivas and Minton (2008). Macromolecular crowding. Annu Rev Biophys "
    "37:375. Scaled-particle excluded-volume theory. [full]",
    "Minton (2001). The influence of macromolecular crowding. J Biol Chem "
    "276:10577. Activity coefficients from scaled-particle theory. [2]",
    "Sasaki et al. (2007). Regulation of DNA nucleases by molecular crowding. "
    "Nucleic Acids Res 35:4086. Exonuclease turnover down about 5.5 times at 20 "
    "percent PEG 8000. [full]",
    "Sasaki et al. (2006). Crowding and DNA polymerase. Biotechnol J 1:440. "
    "Polymerase binding up, intrinsic turnover down. [abs]",
    "Akabayov et al. (2013). Macromolecular crowding and DNA replication. Nat "
    "Commun 4:1615. Replication up about 2 to 4 times at 4 percent PEG. [full]",
    "Zimmerman and Harrison (1987). Macromolecular crowding increases polymerase "
    "binding to DNA. PNAS 84:1871. Enzyme binding enhancement. [abs]",
    "Baltierra-Jasso et al. (2015). Crowding-induced hybridization of single DNA "
    "hairpins. JACS 137:16020. PEG 8000 closing rate up about 4 times. [full]",
    "Homchaudhuri et al. (2006). Effect of crowding by dextrans and Ficolls. "
    "Biopolymers 83:477. Ficoll far milder than extended polymers. [abs]",
]
for r_ in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    add_rich(p, r_, size=8.5)

doc.add_paragraph()
n = doc.add_paragraph()
add_rich(n, "Implementation: crowded_mechanistic.py solves the seven-species "
            "model; the interactive page (crowding_explorer.html) runs the same "
            "model live. The two-variable form is the optional simplification of "
            "Section 10.", italic=True, size=9, color=GREY)

# =============================================================================
# SAVE
# =============================================================================
out = os.path.join(HERE, "Crowded_PredatorPrey_Mechanistic_Model.docx")
try:
    doc.save(out)
except PermissionError:
    base, ext = os.path.splitext(out)
    i = 2
    while True:
        alt = f"{base}_v{i}{ext}"
        try:
            doc.save(alt)
            out = alt
            break
        except PermissionError:
            i += 1
print(f"saved {out}")
